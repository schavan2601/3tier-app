from flask import Flask, request, jsonify
import mysql.connector
from docx import Document
from openpyxl import Workbook
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import os

app = Flask(__name__)

# Database configuration
db_config = {
    'host': 'database',
    'user': 'root',
    'password': 'secret',
    'database': 'appdb'
}

@app.route('/process', methods=['POST'])
def process():
    name = request.form['name']
    email = request.form['email']

    # Save to database
    conn = mysql.connector.connect(**db_config)
    cursor = conn.cursor()
    cursor.execute("INSERT INTO users (name, email) VALUES (%s, %s)", (name, email))
    conn.commit()
    cursor.close()
    conn.close()

    # Generate files
    generate_html(name, email)
    generate_docx(name, email)
    generate_pdf(name, email)
    generate_excel(name, email)

    return "Data processed and saved!"

def generate_html(name, email):
    with open('/app/shared/output.html', 'w') as f:
        f.write(f"<h1>User Info</h1><p>Name: {name}</p><p>Email: {email}</p>")

def generate_docx(name, email):
    doc = Document()
    doc.add_heading('User Info', 0)
    doc.add_paragraph(f"Name: {name}")
    doc.add_paragraph(f"Email: {email}")
    doc.save('/app/shared/output.docx')

def generate_pdf(name, email):
    c = canvas.Canvas('/app/shared/output.pdf', pagesize=letter)
    c.drawString(100, 750, "User Info")
    c.drawString(100, 730, f"Name: {name}")
    c.drawString(100, 710, f"Email: {email}")
    c.save()

def generate_excel(name, email):
    wb = Workbook()
    ws = wb.active
    ws.title = "User Info"
    ws['A1'] = "Name"
    ws['B1'] = "Email"
    ws['A2'] = name
    ws['B2'] = email
    wb.save('/app/shared/output.xlsx')

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
    